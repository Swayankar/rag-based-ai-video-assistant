import os
import io
import uuid
import shutil
import traceback
import threading
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse
from dotenv import load_dotenv

from utils.audio_processor import process_input
from core.transcriber import transcribe_all
from core.summarizer import summarize, generate_title
from core.extractor import extract_action_items, extract_key_decisions, extract_questions
from core.RAG_engine import build_rag_chain, ask_question

# Temp imports
import ssl
import certifi
import requests
import traceback

load_dotenv()

app = FastAPI(title="AI Video Assistant API")

# ALLOWED_ORIGINS = [
#     "http://localhost:5173",
#     "http://127.0.0.1:5173",
#     os.getenv("FRONTEND_ORIGIN", "*"),
# ]

frontend_origin = os.getenv("FRONTEND_ORIGIN")

ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
]

if frontend_origin:
    ALLOWED_ORIGINS.append(frontend_origin)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "downloads"
EXPORT_DIR = "generated"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

jobs: dict = {}
rag_chains: dict = {}

STEPS = [
    ("upload", "Preparing your media"),
    ("audio", "Extracting audio"),
    ("transcribe", "Transcribing speech"),
    ("title", "Naming the meeting"),
    ("summary", "Summarizing the discussion"),
    ("actions", "Pulling out action items"),
    ("decisions", "Finding key decisions"),
    ("questions", "Collecting open questions"),
    ("index", "Indexing for chat"),
    ("done", "All done"),
]
STEP_KEYS = [s[0] for s in STEPS]

def _new_job() -> dict:
    return {
        "status": "queued",       # queued | running | done | error
        "step": STEP_KEYS[0],
        "step_label": STEPS[0][1],
        "progress": 0,
        "error": None,
        "result": None,
        "created_at": datetime.utcnow().isoformat(),
    }

def _set_step(job_id: str, step_key: str):
    idx = STEP_KEYS.index(step_key)
    label = STEPS[idx][1]
    jobs[job_id].update(
        status="running",
        step=step_key,
        step_label=label,
        progress=round((idx / (len(STEPS) - 1)) * 100),
    )

def _run_pipeline(job_id: str, source: str, language: str):
    try:
        _set_step(job_id, "upload")
        chunks = process_input(source)

        _set_step(job_id, "audio")

        _set_step(job_id, "transcribe")
        transcript = transcribe_all(chunks, language=language)

        _set_step(job_id, "title")
        title = generate_title(transcript)

        _set_step(job_id, "summary")
        summary = summarize(transcript)

        _set_step(job_id, "actions")
        action_items = extract_action_items(transcript)

        _set_step(job_id, "decisions")
        decisions = extract_key_decisions(transcript)

        _set_step(job_id, "questions")
        questions = extract_questions(transcript)

        _set_step(job_id, "index")
        rag_chain = build_rag_chain(transcript)
        rag_chains[job_id] = rag_chain

        result = {
            "title": title,
            "transcript": transcript,
            "summary": summary,
            "action_items": action_items,
            "key_decisions": decisions,
            "open_questions": questions,
        }

        _set_step(job_id, "done")
        jobs[job_id].update(status="done", progress=100, result=result)

    except Exception as exc:  # noqa: BLE001
        traceback.print_exc()
        jobs[job_id].update(status="error", error=str(exc))


@app.get("/api/health")
def health():
    return {"ok": True}


@app.post("/api/jobs")
async def create_job(
    language: str = Form("english"),
    youtube_url: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
):
    if not youtube_url and not file:
        raise HTTPException(400, "Provide either a youtube_url or a file.")

    job_id = uuid.uuid4().hex[:12]
    jobs[job_id] = _new_job()

    if file is not None:
        safe_name = f"{job_id}_{file.filename}"
        dest_path = os.path.join(UPLOAD_DIR, safe_name)
        with open(dest_path, "wb") as out:
            shutil.copyfileobj(file.file, out)
        source = dest_path
    else:
        source = youtube_url.strip()

    thread = threading.Thread(
        target=_run_pipeline, args=(job_id, source, language.lower()), daemon=True
    )
    thread.start()

    return {"job_id": job_id}

@app.get("/api/jobs/{job_id}")
def get_job(job_id: str):
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    return job

@app.get("/api/jobs/{job_id}/stream")
async def stream_job(job_id: str):
    """Server-Sent Events feed of job progress, ~2 updates/sec."""
    import asyncio
    import json

    async def event_gen():
        last_payload = None
        while True:
            job = jobs.get(job_id)
            if not job:
                yield f"data: {json.dumps({'status': 'error', 'error': 'not found'})}\n\n"
                return
            payload = json.dumps(job)
            if payload != last_payload:
                yield f"data: {payload}\n\n"
                last_payload = payload
            if job["status"] in ("done", "error"):
                return
            await asyncio.sleep(0.5)

    return StreamingResponse(event_gen(), media_type="text/event-stream")

@app.post("/api/jobs/{job_id}/chat")
async def chat(job_id: str, payload: dict):
    question = (payload or {}).get("question", "").strip()
    if not question:
        raise HTTPException(400, "question is required")

    rag_chain = rag_chains.get(job_id)
    if rag_chain is None:
        raise HTTPException(409, "This meeting isn't indexed for chat yet.")

    answer = ask_question(rag_chain, question)
    return {"answer": answer}

def _plain_join(*parts: str) -> str:
    return "\n\n".join(p for p in parts if p)

@app.get("/api/jobs/{job_id}/download")
def download(job_id: str, fmt: str = "pdf"):
    job = jobs.get(job_id)
    if not job or job["status"] != "done":
        raise HTTPException(409, "Result not ready yet.")
    result = job["result"]

    if fmt == "docx":
        from docx import Document

        doc = Document()
        doc.add_heading(result["title"] or "Meeting Notes", level=0)
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(result["summary"])
        doc.add_heading("Action Items", level=1)
        doc.add_paragraph(result["action_items"])
        doc.add_heading("Key Decisions", level=1)
        doc.add_paragraph(result["key_decisions"])
        doc.add_heading("Open Questions", level=1)
        doc.add_paragraph(result["open_questions"])
        doc.add_heading("Full Transcript", level=1)
        doc.add_paragraph(result["transcript"])

        out_path = os.path.join(EXPORT_DIR, f"{job_id}.docx")
        doc.save(out_path)
        return FileResponse(
            out_path,
            filename=f"{(result['title'] or 'meeting-notes').strip()[:50]}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    elif fmt == "pdf":
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch

        out_path = os.path.join(EXPORT_DIR, f"{job_id}.pdf")
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("h1", parent=styles["Heading1"])
        h2 = ParagraphStyle("h2", parent=styles["Heading2"])
        body = ParagraphStyle("body", parent=styles["BodyText"], spaceAfter=10)

        doc = SimpleDocTemplate(out_path, pagesize=LETTER,
                                 topMargin=0.8 * inch, bottomMargin=0.8 * inch)
        story = [Paragraph(result["title"] or "Meeting Notes", h1), Spacer(1, 12)]

        def section(title, text):
            story.append(Paragraph(title, h2))
            for line in (text or "").split("\n"):
                if line.strip():
                    story.append(Paragraph(line, body))
            story.append(Spacer(1, 8))

        section("Summary", result["summary"])
        section("Action Items", result["action_items"])
        section("Key Decisions", result["key_decisions"])
        section("Open Questions", result["open_questions"])
        section("Full Transcript", result["transcript"])

        doc.build(story)
        return FileResponse(
            out_path,
            filename=f"{(result['title'] or 'meeting-notes').strip()[:50]}.pdf",
            media_type="application/pdf",
        )

    raise HTTPException(400, "fmt must be 'pdf' or 'docx'")

# Temparary endpoints to check the HuggingFace issue

@app.get("/api/test")
def test():
    r = requests.get("https://www.google.com", timeout=10)
    return {"status": r.status_code}

@app.get("/api/test-youtube")
def test_youtube():
    try:
        r = requests.get("https://www.youtube.com", timeout=10)
        return {
            "status": r.status_code,
            "headers": dict(r.headers),
        }
    except Exception as e:
        return {
            "error": type(e).__name__,
            "message": str(e),
            "traceback": traceback.format_exc(),
        }

@app.get("/api/debug")
def debug():
    return {
        "python": os.sys.version,
        "openssl": ssl.OPENSSL_VERSION,
        "requests": requests.__version__,
        "certifi": certifi.__version__,
        "ca_bundle": certifi.where(),
    }